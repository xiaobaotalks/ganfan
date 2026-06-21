"""生成干饭省省图文发布 docx 文件"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

BASE_DIR = r'D:\C-AI\m-code\m-qwen\logs'
OUTPUT_FILE = r'D:\C-AI\m-code\m-qwen\docs\干饭省省_发布图文.docx'

XHS_IMAGES = [
    ('v18_home.png', '图1 - 首页：暖橙主题 + 1 屏看懂全部功能'),
    ('v3_capture.png', '图2 - 拍照页：Qwen3-VL-2B 端侧推理，招牌照片 100% 本地处理'),
    ('v10_i.png', '图3 - 结果页：1 张图 = 5 平台价格 + 附近 12 家同款门店'),
    ('v9_real3.png', '图4 - 联网模式：半真实数据策略，5 平台 deeplink 直跳'),
    ('v20_flash.png', '图5 - 秒省页：每天 12/18 点限时秒杀'),
    ('v16_circle.png', '图6 - 省圈页：5 个本地干饭群，拼单拼团更便宜'),
    ('v16_me.png', '图7 - 我的页：模式开关 + 4 个对话框，隐私透明可查'),
    ('v7_history_full.png', '图8 - 识别记录：真实街拍记录，省多少钱一目了然'),
]

XHS_TITLE = '端侧 AI 干饭省省 · 1 秒扫招牌,5 平台比价'

XHS_TAGS = [
    '干饭省省', '端侧AI', 'Qwen3VL', 'MNN', '拍招牌识餐厅',
    '团购比价', '隐私不上传', '离线可用', 'AI工具', '大模型应用',
    'iQOONeo11', '学生党福利', '吃饭省钱'
]

XHS_BODY = """姐妹们兄弟们，出门吃饭的最大痛点，我做了 1 个 App 解决了 🫶

🎯 项目出发点（为什么做这个）

以前想去一家店，得打开 5 个 App（抖音团购 / 美团团购 / 美团外卖 / 淘宝闪购 / 京东秒送）挨个查，比价 5 分钟，最后多花 30% 😭

更糟的是，很多人其实就在店里点餐了，压根不知道隔壁同一品牌另一家店有 ¥8 团购。信息差 = 钱差。

我做了个「干饭省省」：举起手机扫招牌，1 秒看到 5 平台最低价 + 附近 12 家同款门店 ⬇️

🌟 项目特点（端侧 AI 创新挑战赛作品 · 4 大亮点）

1. 🔥 Qwen3-VL-2B 大模型跑在手机上（MNN 端侧推理框架）— 不是云端 OCR，是真正能"看招牌"的多模态大模型
2. 🔥 完全离线 — 招牌照片不出手机，飞行模式也能用，飞行模式实测 13.86 tok/s
3. 🔥 半真实数据策略 — 招牌本地识别 → 高德 POI 拿真实店 → 5 平台价格走 deeplink 跳官方 App 实时查
4. 🔥 按 GPS 选最近分店 — 7 品牌 × 3-4 分店，招牌"luckin"识别后自动选离你最近的瑞幸（实测 101m）

🎮 使用体验（操作 5 步 · 1 顿饭省 ¥5-30）

1. 📍 看到想吃的店
2. 🏠 打开 App，点大圆拍照
3. 📸 对准招牌，1 拍
4. ⚡ 1-2 秒看到 5 平台比价（团购 2 + 外卖 3）+ 附近 12 家同款门店（2km 内）
5. 🎟 点最低价卡片 → 直接跳到对应 App 买券

🏆 实战战绩（我朋友实测）：
- 瑞幸咖啡 五道口店 101m → 抖音团购 ¥9（比星巴克最低 ¥30 省 ¥21）
- 麦当劳 五道口店 → 淘宝闪购麦满分 ¥11（比到店 ¥22.9 省 ¥12）
- 海底捞 4 人套餐 抖音 ¥299（比京东 ¥599 省一半）

单人 9 天极限开发 · Demo 已跑通 · 端侧 AI 创新挑战赛参赛作品"""

XHS_COMMENTS = """你们最常去哪家店？5 平台一般谁最便宜？评论区聊聊 🗣
想知道技术细节（Qwen3-VL 怎么部署 / 性能怎么样）的扣 1，我下条笔记详细拆
想看我实测视频的扣 2，后续更 ⚡
关注我，后续更新技术细节 + 街拍实测 + 省钱技巧 💰"""

WX_TITLES = [
    '我做了一个 App：扫一下餐厅招牌，1 秒比价 5 个平台，全程离线不联网',
    'Qwen3-VL-2B 跑在手机上：9 天单人开发，一个端侧 AI 干饭比价 App 的诞生',
    '每次吃饭多花 30%？我用端侧 AI 做了个「扫招牌比价」工具',
]

WX_BODY = """大家好，我是干饭省省的开发者。

今天想和大家聊聊我最近花了 9 天时间，一个人从零做出来的一款端侧 AI 应用 —— 干饭省省。

这不是什么高大上的企业级项目，只是一个解决我自己痛点的小东西。但做出来之后发现，它可能也是很多人的痛点。

---

## 一、为什么做这个？

不知道你有没有过这种经历：

走到一家餐厅门口，想看看有没有团购优惠，于是打开抖音搜一下；
搜完觉得可能美团更便宜，又切到美团；
想起淘宝闪购可能有新人价，再开淘宝；
最后京东秒送好像也有补贴……

来来回回切换五六个 App，比价 5 分钟，最后可能还是选错了，多花了冤枉钱。

更气人的是，很多时候你就站在店里点餐了，根本不知道同一条街 100 米外的同品牌另一家店，有便宜一半的团购券。

**信息差，就是钱差。**

于是我想：能不能做一个东西，**举起手机扫一下招牌，立刻告诉我这家店在所有平台的最低价，以及附近同款门店的价格？**

说干就干，9 天时间，我做出来了。

---

## 二、干饭省省是什么？

简单说，就是一个「**拍招牌，比低价**」的 App。

对准餐厅招牌拍一张照，1-2 秒后，你就能看到：

- ✅ **5 个平台的实时价格**（抖音团购、美团团购、美团外卖、淘宝闪购、京东秒送）
- ✅ **附近 12 家同品牌门店**（2km 范围内，按距离排序）
- ✅ **一键跳转到对应 App 买券**（deeplink 直达，不用再搜）

整个过程，**招牌照片不需要上传到云端**，全部在手机本地处理。

---

## 三、技术亮点：Qwen3-VL 大模型跑在手机上

这可能是整个项目最有意思的地方。

### 1. 真正的端侧 VLM，不是云端 OCR

很多人以为这就是个拍照识别文字的 OCR 工具，其实不是。

我用的是 **Qwen3-VL-2B** 多模态大模型，通过 **MNN 端侧推理框架** 部署到手机上。INT4 量化后，模型可以在手机上流畅运行，飞行模式下实测 **13.86 tok/s**。

为什么不用 OCR？因为 OCR 只能识别文字，而 VLM 可以**理解招牌的含义**：

- 识别品牌 logo 而不仅是文字
- 理解餐厅类型（咖啡 / 快餐 / 火锅等）
- 抗干扰能力更强（角度倾斜、光线不足、部分遮挡）

### 2. 完全离线，隐私不上传

这是我最在意的一点。

你拍的餐厅招牌照片，**100% 在手机本地处理**，不会上传到任何服务器。哪怕开飞行模式，App 也能用。

在「我的」页面有个模式开关，默认是完全离线模式。

### 3. 半真实数据策略

当然，完全离线也有局限 —— 你没法拿到实时的团购价格。

所以我设计了一套「半真实数据策略」：

- **招牌识别**：本地 VLM 完成（离线）
- **门店信息**：联网时通过高德 POI 拉取真实门店数据
- **价格查询**：通过 deeplink 跳转到官方 App 查看实时价格

这样既保证了核心功能（招牌识别）的隐私性，又能获取真实的门店和价格信息。

### 4. GPS 自动选最近分店

同一个品牌，一条街可能有好几家店。

干饭省省会根据你的 GPS 位置，自动匹配最近的分店。比如识别到 "luckin"，会自动定位到离你 101 米的那家瑞幸，而不是 1.8 公里外的。

---

## 四、除了比价，还有这些功能

### ⚡ 秒省页 —— 限时秒杀

每天 12:00 和 18:00 两个饭点，精选 4 款超值秒杀商品，一键直达。

### 👥 省圈页 —— 找饭搭子

加入本地干饭群，和街坊邻居一起拼单拼团，薅更多羊毛。

### 📋 识别记录

所有扫过的招牌都会保留历史记录，每顿饭省了多少钱一目了然。

---

## 五、实战战绩

说了这么多，到底能省多少钱？

我朋友实测了几次：

| 餐厅 | 平台 | 价格 | 对比 | 省了多少 |
|------|------|------|------|----------|
| 瑞幸咖啡（五道口店 101m） | 抖音团购 | ¥9 | 星巴克最低 ¥30 | ¥21 |
| 麦当劳（五道口店） | 淘宝闪购 | ¥11 | 到店 ¥22.9 | ¥12 |
| 海底捞（4 人套餐） | 抖音 | ¥299 | 京东 ¥599 | ¥300 |

一顿饭省个几块到几十块不等，积少成多，一个月下来也能省出不少奶茶钱。

---

## 六、写在最后

这是我参加「端侧 AI 创新挑战赛」的作品，单人 9 天极限开发完成。

做这个项目的初衷，是想探索一下端侧大模型到底能做什么落地的应用。现在看来，**端侧 AI 不只是技术炫技，它真的能解决实际问题，而且能保护隐私**。

如果你对技术细节感兴趣 —— Qwen3-VL 怎么部署到 MNN、INT4 量化怎么做、性能调优有哪些坑，欢迎点赞 + 在看，人多的话我下篇详细拆。

也欢迎在评论区聊聊：你平时吃饭会比价吗？一般哪个平台最划算？

---

**本文是端侧 AI 创新挑战赛参赛作品，项目已开源，欢迎交流。**

#干饭省省 #端侧AI #Qwen3VL #MNN #大模型应用"""


def add_styled_heading(doc, text, level=1, color=None):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading


def add_body_text(doc, text):
    """添加正文段落"""
    lines = text.split('\n')
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
    return p


def main():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    
    # 文档标题
    title = doc.add_heading('干饭省省 · 发布图文素材', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph('本文档包含小红书和微信公众号两个平台的发布素材，可直接复制使用。')
    doc.add_paragraph()
    
    # ========== 小红书部分 ==========
    doc.add_heading('一、小红书发布素材', level=1)
    
    # 标题
    doc.add_heading('1.1 标题', level=2)
    p = doc.add_paragraph()
    run = p.add_run(XHS_TITLE)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(255, 107, 107)
    
    # 配图
    doc.add_heading('1.2 配图（共 8 张）', level=2)
    doc.add_paragraph('建议按以下顺序上传：')
    
    for i, (img_file, desc) in enumerate(XHS_IMAGES):
        img_path = os.path.join(BASE_DIR, img_file)
        if os.path.exists(img_path):
            doc.add_paragraph(f'第 {i+1} 张：{desc}', style='List Number')
            doc.add_picture(img_path, width=Inches(2.5))
            doc.add_paragraph()
        else:
            doc.add_paragraph(f'第 {i+1} 张：{desc}（文件不存在：{img_file}）')
    
    # 正文
    doc.add_heading('1.3 正文内容', level=2)
    add_body_text(doc, XHS_BODY)
    
    # 标签
    doc.add_heading('1.4 标签', level=2)
    tags_str = ' '.join([f'#{tag}' for tag in XHS_TAGS])
    p = doc.add_paragraph()
    run = p.add_run(tags_str)
    run.font.color.rgb = RGBColor(0, 122, 255)
    
    # 评论引导
    doc.add_heading('1.5 评论区引导', level=2)
    add_body_text(doc, XHS_COMMENTS)
    
    # 发布建议
    doc.add_heading('1.6 发布建议', level=2)
    doc.add_paragraph('• 发布时间：工作日 19:00-21:00 / 周末 11:00-13:00 / 17:00-19:00')
    doc.add_paragraph('• @账号：@阿里达摩院MNN @通义千问Qwen @vivo')
    doc.add_paragraph('• 评论区每条必回 + 带 emoji + 引导关注')
    
    doc.add_page_break()
    
    # ========== 公众号部分 ==========
    doc.add_heading('二、微信公众号发布素材', level=1)
    
    # 标题
    doc.add_heading('2.1 标题（3 个备选）', level=2)
    for i, title in enumerate(WX_TITLES):
        p = doc.add_paragraph()
        run = p.add_run(f'备选 {i+1}：{title}')
        if i == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 107, 107)
    
    # 封面图
    doc.add_heading('2.2 封面图', level=2)
    doc.add_paragraph('主图：v18_home.png（暖橙首页，功能丰富）')
    img_path = os.path.join(BASE_DIR, 'v18_home.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(4.0))
    
    doc.add_paragraph()
    doc.add_paragraph('次图（分享小图）：v10_i.png（结果页，5 平台比价最有冲击力）')
    img_path2 = os.path.join(BASE_DIR, 'v10_i.png')
    if os.path.exists(img_path2):
        doc.add_picture(img_path2, width=Inches(2.5))
    
    # 正文
    doc.add_heading('2.3 正文内容', level=2)
    
    # 分段处理正文
    sections = WX_BODY.split('---')
    for section in sections:
        section = section.strip()
        if not section:
            continue
        
        lines = section.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            if line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=3)
            elif line.startswith('|') and '---' not in line:
                # 表格行
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.replace('**', ''))
                run.bold = True
            elif line.startswith('#') and not line.startswith('##'):
                # 标签
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.color.rgb = RGBColor(0, 122, 255)
            else:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = '微软雅黑'
    
    # 发布建议
    doc.add_heading('2.4 发布建议', level=2)
    doc.add_paragraph('• 发布时间：周二 / 周四 20:00-21:30（技术类阅读高峰）')
    doc.add_paragraph('• 开头引导：加一句「星标我的公众号，第一时间看端侧 AI 实战」')
    doc.add_paragraph('• 结尾互动：投票「你最想看哪类内容？技术拆解 / 实测视频 / 更多功能」')
    doc.add_paragraph('• 原文链接：可以放项目 GitHub 仓库（如果有的话）')
    doc.add_paragraph('• 排版建议：重点内容用加粗或引用块，图片之间加分割线')
    
    # 保存
    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f'✅ 文档已生成：{OUTPUT_FILE}')


if __name__ == '__main__':
    main()
